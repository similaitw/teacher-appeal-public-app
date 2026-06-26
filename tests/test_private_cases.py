from __future__ import annotations

import sqlite3
import sys
from pathlib import Path
import csv
import json

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

import private_db
import private_documents
from private_db import (
    ParsedUnit,
    create_case,
    export_analysis_package,
    init_private_db,
    insert_document_with_units,
    safe_filename,
    search_units,
)
from private_documents import import_document, import_document_from_path


@pytest.fixture()
def isolated_private_store(tmp_path, monkeypatch):
    db_path = tmp_path / "private_cases.db"
    upload_dir = tmp_path / "uploaded_cases"
    export_dir = tmp_path / "exports"
    monkeypatch.setattr(private_db, "PRIVATE_DB_PATH", db_path)
    monkeypatch.setattr(private_db, "UPLOADED_CASES_DIR", upload_dir)
    monkeypatch.setattr(private_db, "PRIVATE_EXPORTS_DIR", export_dir)
    monkeypatch.setattr(private_documents, "UPLOADED_CASES_DIR", upload_dir)
    init_private_db(db_path)
    return db_path, upload_dir, export_dir


def make_case(db_path: Path) -> dict[str, object]:
    return create_case("A-001", "測試案件", "申訴", "已去識別化測試", db_path=db_path)


def add_document(db_path: Path, case_id: int, filename: str, units: list[ParsedUnit], mime_type: str = "") -> dict[str, object]:
    doc, duplicate = insert_document_with_units(
        case_id=case_id,
        original_filename=filename,
        stored_filename=filename,
        mime_type=mime_type,
        file_sha256=f"sha-{filename}",
        file_size=123,
        parse_status="parsed",
        parse_error="",
        units=units,
        db_path=db_path,
    )
    assert not duplicate
    return doc


def test_txt_import_keeps_line_numbers(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    text = "第一行事實\n第二行理由\n第三行結論\n"
    result = import_document(int(case["id"]), "report.txt", text.encode("utf-8"), "text/plain", db_path=db_path)
    assert result["parse_status"] == "parsed"
    rows = search_units("理由", case_id=int(case["id"]), db_path=db_path)
    assert rows
    assert rows[0]["line_start"] == 1
    assert rows[0]["line_end"] == 3


def test_docx_import_keeps_paragraph_numbers(isolated_private_store, tmp_path):
    from docx import Document

    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    path = tmp_path / "appeal.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段 教師申訴")
    doc.save(path)
    result = import_document_from_path(int(case["id"]), path, db_path=db_path)
    assert result["parse_status"] == "parsed"
    rows = search_units("教師申訴", case_id=int(case["id"]), db_path=db_path)
    assert rows[0]["paragraph_number"] == 2


def test_text_layer_pdf_import_keeps_page_numbers(isolated_private_store, tmp_path):
    import fitz

    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    path = tmp_path / "appeal.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "page one appeal")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "page two reasoning")
    pdf.save(path)
    pdf.close()
    result = import_document_from_path(int(case["id"]), path, db_path=db_path)
    assert result["parse_status"] == "parsed"
    rows = search_units("reasoning", case_id=int(case["id"]), db_path=db_path)
    assert rows[0]["page_number"] == 2


def test_duplicate_file_detection(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    data = "相同內容".encode("utf-8")
    first = import_document(int(case["id"]), "same.txt", data, "text/plain", db_path=db_path)
    second = import_document(int(case["id"]), "same-copy.txt", data, "text/plain", db_path=db_path)
    assert not first["duplicate"]
    assert second["duplicate"]
    with sqlite3.connect(db_path) as conn:
        assert conn.execute("SELECT COUNT(*) FROM private_documents").fetchone()[0] == 1


def test_safe_filename_blocks_path_traversal():
    assert safe_filename("../../危險/../case?.txt") == "case_.txt"
    assert safe_filename("..\\..\\evil.pdf").endswith("evil.pdf")


def test_single_case_search_is_isolated(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case_a = create_case("A", "A", "", "", db_path=db_path)
    case_b = create_case("B", "B", "", "", db_path=db_path)
    import_document(int(case_a["id"]), "a.txt", "共同詞 A案".encode("utf-8"), "text/plain", db_path=db_path)
    import_document(int(case_b["id"]), "b.txt", "共同詞 B案".encode("utf-8"), "text/plain", db_path=db_path)
    rows_a = search_units("共同詞", case_id=int(case_a["id"]), db_path=db_path)
    assert rows_a
    assert {row["case_uuid"] for row in rows_a} == {case_a["case_uuid"]}
    rows_all = search_units("共同詞", case_id=None, db_path=db_path)
    assert {row["case_uuid"] for row in rows_all} == {case_a["case_uuid"], case_b["case_uuid"]}


def test_markdown_analysis_package_source_labels(isolated_private_store):
    db_path, _, export_dir = isolated_private_store
    case = make_case(db_path)
    import_document(int(case["id"]), "source.txt", "第1行\n第2行\n".encode("utf-8"), "text/plain", db_path=db_path)
    out_dir = export_analysis_package(int(case["id"]), db_path=db_path)
    assert out_dir == export_dir / str(case["case_uuid"])
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert "# 案件基本資料" in full_context
    assert "[來源：D001，第1至2行]" in full_context
    assert (out_dir / "case_manifest.json").exists()
    assert (out_dir / "source_index.csv").exists()


def test_full_context_export_writes_required_files_and_manifest(isolated_private_store):
    db_path, _, export_dir = isolated_private_store
    case = make_case(db_path)
    add_document(db_path, int(case["id"]), "full.txt", [ParsedUnit("txt_lines", "原文第一行\n原文第二行", line_start=1, line_end=2)])
    out_dir = export_analysis_package(int(case["id"]), db_path=db_path)
    assert out_dir == export_dir / str(case["case_uuid"])
    assert sorted(p.name for p in out_dir.iterdir()) == ["case_full_context.md", "case_manifest.json", "source_index.csv"]
    manifest = json.loads((out_dir / "case_manifest.json").read_text(encoding="utf-8"))
    assert manifest["case_uuid"] == case["case_uuid"]
    assert manifest["case_number"] == "A-001"
    assert manifest["title"] == "測試案件"
    assert manifest["case_type"] == "申訴"
    assert manifest["document_count"] == 1
    assert manifest["unit_count"] == 1
    assert manifest["source_files"][0]["filename"] == "full.txt"
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert "# 案件基本資料" in full_context
    assert "| 文件代號 | 原始檔名 | 文件類型 | 頁數或段落數 |" in full_context
    assert "原文第一行\n原文第二行" in full_context


def test_multi_document_export_sorts_by_document_id(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    second = add_document(db_path, int(case["id"]), "b.txt", [ParsedUnit("txt_lines", "B", line_start=1, line_end=1)])
    first = add_document(db_path, int(case["id"]), "a.txt", [ParsedUnit("txt_lines", "A", line_start=1, line_end=1)])
    out_dir = export_analysis_package(int(case["id"]), document_ids=[int(first["id"]), int(second["id"])], db_path=db_path)
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert full_context.index("## 文件 D001：b.txt") < full_context.index("## 文件 D002：a.txt")


def test_pdf_page_source_and_empty_page_marker(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    add_document(
        db_path,
        int(case["id"]),
        "pages.pdf",
        [
            ParsedUnit("pdf_page", "第一頁原文", page_number=1),
            ParsedUnit("pdf_page", "", page_number=2),
        ],
        "application/pdf",
    )
    out_dir = export_analysis_package(int(case["id"]), db_path=db_path)
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert "### 第 1 頁" in full_context
    assert "[來源：D001，第1頁]" in full_context
    assert "### 第 2 頁" in full_context
    assert "[來源：D001，第2頁]" in full_context
    assert "[本頁無法擷取文字]" in full_context


def test_docx_paragraph_source(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    add_document(db_path, int(case["id"]), "memo.docx", [ParsedUnit("docx_paragraph", "第十五段原文", paragraph_number=15)])
    out_dir = export_analysis_package(int(case["id"]), db_path=db_path)
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert "### 第 15 段" in full_context
    assert "[來源：D001，第15段]" in full_context
    assert "第十五段原文" in full_context


def test_txt_line_source_and_source_index_columns(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    add_document(db_path, int(case["id"]), "notes.txt", [ParsedUnit("txt_lines", "第20至30行原文", line_start=20, line_end=30)])
    out_dir = export_analysis_package(int(case["id"]), db_path=db_path)
    full_context = (out_dir / "case_full_context.md").read_text(encoding="utf-8")
    assert "### 第 20 至 30 行" in full_context
    assert "[來源：D001，第20至30行]" in full_context
    with (out_dir / "source_index.csv").open("r", encoding="utf-8-sig", newline="") as fh:
        rows = list(csv.DictReader(fh))
    assert rows[0]["source_id"] == "D001"
    assert rows[0]["filename"] == "notes.txt"
    assert rows[0]["unit_type"] == "txt_lines"
    assert rows[0]["line_start"] == "20"
    assert rows[0]["line_end"] == "30"


def test_empty_case_export_raises_clear_error(isolated_private_store):
    db_path, _, _ = isolated_private_store
    case = make_case(db_path)
    with pytest.raises(ValueError, match="尚未匯入任何文件"):
        export_analysis_package(int(case["id"]), db_path=db_path)


def test_public_database_not_modified_by_private_schema(isolated_private_store):
    public_db = Path(__file__).resolve().parents[1] / "data" / "appeal_cases.db"
    if not public_db.exists():
        pytest.skip("公開資料庫不存在")
    with sqlite3.connect(public_db) as conn:
        names = {row[0] for row in conn.execute("SELECT name FROM sqlite_master WHERE type = 'table'")}
    assert "cases" in names
    assert "cases_fts" in names
    assert "private_cases" not in names
    assert "private_document_units_fts" not in names
